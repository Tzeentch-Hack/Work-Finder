import io
from docx import Document
from docx.shared import Inches
from resume_generator import ResumeBuilder

import subprocess
import os
import time

resume_path = "./CVs"

def compile_to_pdf(input_path):
    subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', input_path])


def build_pdf_resume(resume_info, image_path, username):
    path = build_docx_resume(resume_info, image_path)
    compile_to_pdf(path)

def build_docx_resume(resume_info, image_path, username):
    doc = Document()
    doc.add_picture(image_path, width=Inches(1.5))
    p = doc.add_paragraph()
    p.add_run(resume_info['contact_information']).bold = True


    doc.add_heading('Professional Profile', level=1)
    doc.add_paragraph(resume_info['professional_profile'])

    doc.add_heading('Skills', level=1)
    doc.add_paragraph(resume_info['skills'])

    doc.add_heading('Additional Information', level=1)
    doc.add_paragraph(resume_info['additional_information'])

    #doc.save('temp_resume.docx')
    return save_to_user_folder(doc, username)


def save_to_user_folder(doc, username):
    if not os.path.exists(resume_path):
        os.makedirs(resume_path)
    user_path = os.path.join(resume_path, username)
    if not os.path.exists(user_path):
        os.makedirs(user_path)
    try:
        current_time_seconds = int(time.time())
        time_string = str(current_time_seconds)
        path = os.path.join(user_path, f'{username}_cv_{time_string}.docx')
        doc.save(path)
        return path
    except Exception("IO error") as e:
        raise e


if __name__ == "__main__":
    resume_builder = ResumeBuilder("gpt-3.5-turbo", 2000, 'ru')
    user_info = {
        'age': 30,
        'gender': 'Male',
        'education': 'No education',
        'place_of_living': 'Tashkent',
        'type_of_work': 'Physical',
        'work_schedule': 'Full-time',
        'form_of_work': 'Local',
        'specialization': 'Cleaner',
        'email': 'example@email.com',
        'phone': '+1 123-456-7890'
    }
    import time
    start = time.time()
    resume_info = resume_builder.build_resume(**user_info)
    end = time.time()
    print('end - start', end - start)
    image_path = 'igore.jpg'
    build_pdf_resume(resume_info, image_path)
